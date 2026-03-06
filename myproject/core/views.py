from io import BytesIO

from django.contrib import messages
from django.contrib.auth import get_user_model, login
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator
from django.db.models import Count
from django.http import FileResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.utils import timezone

from .forms import (
    CommunityCommentForm,
    CommunityPostForm,
    CouncilJoinApplicationForm,
    EventRegistrationForm,
    FeedbackForm,
    ProfileUpdateForm,
    SignUpForm,
    UserFileUploadForm,
    UserUpdateForm,
)
from .models import (
    CommunityPost,
    CouncilJoinApplication,
    Document,
    Event,
    EventRegistration,
    FAQ,
    FeedbackMessage,
    News,
    Project,
    StudentCouncilMember,
    UserFile,
    UserProfile,
)
from .permissions import has_any_role, role_required

User = get_user_model()


def home(request):
    latest_news = News.objects.filter(is_published=True)[:6]
    upcoming_events = Event.objects.filter(is_published=True, start_at__gte=timezone.now())[:6]
    active_projects = Project.objects.filter(is_published=True)[:6]
    community_posts = CommunityPost.objects.filter(is_published=True).select_related("author")[:3]
    return render(
        request,
        "core/home.html",
        {
            "latest_news": latest_news,
            "upcoming_events": upcoming_events,
            "active_projects": active_projects,
            "community_posts": community_posts,
        },
    )


def news_list(request):
    queryset = News.objects.filter(is_published=True)
    paginator = Paginator(queryset, 10)
    page_obj = paginator.get_page(request.GET.get("page"))
    return render(request, "core/news_list.html", {"page_obj": page_obj})


def news_detail(request, pk):
    item = get_object_or_404(News, pk=pk, is_published=True)
    return render(request, "core/news_detail.html", {"item": item})


def events_list(request):
    queryset = Event.objects.filter(is_published=True)
    archive = request.GET.get("archive")
    if archive == "1":
        queryset = queryset.filter(start_at__lt=timezone.now()).order_by("-start_at")
    else:
        queryset = queryset.filter(start_at__gte=timezone.now()).order_by("start_at")
    return render(request, "core/events_list.html", {"events": queryset, "archive": archive == "1"})


def event_detail(request, pk):
    event = get_object_or_404(Event, pk=pk, is_published=True)
    already_registered = False
    form = None
    if request.user.is_authenticated:
        already_registered = EventRegistration.objects.filter(user=request.user, event=event).exists()
        if request.method == "POST" and not already_registered:
            form = EventRegistrationForm(request.POST)
            if form.is_valid():
                registration = form.save(commit=False)
                registration.user = request.user
                registration.event = event
                registration.save()
                messages.success(request, "Заявка на мероприятие отправлена.")
                return redirect("core:event_detail", pk=event.pk)
        else:
            form = EventRegistrationForm()
    return render(
        request,
        "core/event_detail.html",
        {"event": event, "form": form, "already_registered": already_registered},
    )


def council_info(request):
    members = StudentCouncilMember.objects.all()
    return render(request, "core/council_info.html", {"members": members})


def documents_list(request):
    docs = Document.objects.filter(is_published=True)
    return render(request, "core/documents_list.html", {"documents": docs})


def projects_list(request):
    projects = Project.objects.filter(is_published=True)
    return render(request, "core/projects_list.html", {"projects": projects})


def faq_list(request):
    faq_items = FAQ.objects.filter(is_published=True)
    return render(request, "core/faq_list.html", {"faq_items": faq_items})


def feedback_create(request):
    if request.method == "POST":
        form = FeedbackForm(request.POST)
        if form.is_valid():
            item = form.save(commit=False)
            if request.user.is_authenticated:
                item.user = request.user
            item.save()
            messages.success(request, "Ваше обращение отправлено. Спасибо за обратную связь.")
            return redirect("core:feedback")
    else:
        initial = {}
        if request.user.is_authenticated:
            full_name = f"{request.user.first_name} {request.user.last_name}".strip()
            initial = {"name": full_name or request.user.username, "email": request.user.email}
        form = FeedbackForm(initial=initial)
    return render(request, "core/feedback.html", {"form": form})


def join_request_create(request):
    if request.method == "POST":
        form = CouncilJoinApplicationForm(request.POST)
        if form.is_valid():
            item = form.save(commit=False)
            if request.user.is_authenticated:
                item.user = request.user
            item.save()
            messages.success(request, "Заявка на вступление в студсовет отправлена.")
            return redirect("core:join")
    else:
        initial = {}
        if request.user.is_authenticated:
            full_name = f"{request.user.first_name} {request.user.last_name}".strip()
            profile = getattr(request.user, "profile", None)
            initial = {
                "full_name": full_name or request.user.username,
                "email": request.user.email,
                "course": profile.course if profile else "",
            }
        form = CouncilJoinApplicationForm(initial=initial)
    return render(request, "core/join_request.html", {"form": form})


def community_feed(request):
    post_form = CommunityPostForm()
    comment_form = CommunityCommentForm()

    if request.method == "POST":
        if not request.user.is_authenticated:
            messages.error(request, "Для публикации и комментариев выполните вход.")
            return redirect(f"{reverse('login')}?next={reverse('core:community')}")

        if "create_post" in request.POST:
            post_form = CommunityPostForm(request.POST)
            if post_form.is_valid():
                post = post_form.save(commit=False)
                post.author = request.user
                post.save()
                messages.success(request, "Публикация добавлена в сообщество.")
                return redirect(f"{reverse('core:community')}#post-{post.id}")

        if "add_comment" in request.POST:
            post_id = request.POST.get("post_id")
            post = get_object_or_404(CommunityPost, pk=post_id, is_published=True)
            comment_form = CommunityCommentForm(request.POST)
            if comment_form.is_valid():
                comment = comment_form.save(commit=False)
                comment.author = request.user
                comment.post = post
                comment.save()
                messages.success(request, "Комментарий опубликован.")
                return redirect(f"{reverse('core:community')}#post-{post.id}")

    posts = (
        CommunityPost.objects.filter(is_published=True)
        .select_related("author")
        .prefetch_related("comments__author")
    )

    return render(
        request,
        "core/community_feed.html",
        {
            "posts": posts,
            "post_form": post_form,
            "comment_form": comment_form,
        },
    )


def register(request):
    if request.user.is_authenticated:
        return redirect("core:cabinet")
    if request.method == "POST":
        form = SignUpForm(request.POST)
        if form.is_valid():
            user = form.save()
            profile, _ = UserProfile.objects.get_or_create(user=user)
            profile.role = UserProfile.ROLE_STUDENT
            profile.save(update_fields=["role"])
            login(request, user)
            messages.success(request, "Регистрация выполнена успешно.")
            return redirect("core:cabinet")
    else:
        form = SignUpForm()
    return render(request, "registration/register.html", {"form": form})


@login_required
def cabinet(request):
    feedbacks = FeedbackMessage.objects.filter(user=request.user)[:5]
    registrations = EventRegistration.objects.filter(user=request.user).select_related("event")[:5]
    join_requests = CouncilJoinApplication.objects.filter(user=request.user)[:5]
    my_posts = CommunityPost.objects.filter(author=request.user)[:5]
    my_files = UserFile.objects.filter(owner=request.user)[:5]
    return render(
        request,
        "core/cabinet.html",
        {
            "feedbacks": feedbacks,
            "registrations": registrations,
            "join_requests": join_requests,
            "my_posts": my_posts,
            "my_files": my_files,
        },
    )


@login_required
def profile_edit(request):
    profile, _ = UserProfile.objects.get_or_create(user=request.user)
    if request.method == "POST":
        user_form = UserUpdateForm(request.POST, instance=request.user)
        profile_form = ProfileUpdateForm(request.POST, instance=profile)
        if user_form.is_valid() and profile_form.is_valid():
            user_form.save()
            profile_form.save()
            messages.success(request, "Профиль обновлен.")
            return redirect("core:profile_edit")
    else:
        user_form = UserUpdateForm(instance=request.user)
        profile_form = ProfileUpdateForm(instance=profile)
    return render(
        request,
        "core/profile_edit.html",
        {"user_form": user_form, "profile_form": profile_form},
    )


@login_required
def my_feedbacks(request):
    items = FeedbackMessage.objects.filter(user=request.user)
    return render(request, "core/my_feedbacks.html", {"items": items})


@login_required
def my_events(request):
    items = EventRegistration.objects.filter(user=request.user).select_related("event")
    return render(request, "core/my_events.html", {"items": items})


@login_required
def my_join_requests(request):
    items = CouncilJoinApplication.objects.filter(user=request.user)
    return render(request, "core/my_join_requests.html", {"items": items})


@login_required
def my_posts(request):
    items = CommunityPost.objects.filter(author=request.user).prefetch_related("comments")
    return render(request, "core/my_posts.html", {"items": items})


@login_required
def my_files(request):
    if request.method == "POST":
        form = UserFileUploadForm(request.POST, request.FILES)
        if form.is_valid():
            item = form.save(commit=False)
            item.owner = request.user
            item.save()
            messages.success(request, "Файл загружен в хранилище.")
            return redirect("core:my_files")
    else:
        form = UserFileUploadForm()

    items = UserFile.objects.filter(owner=request.user)
    return render(request, "core/my_files.html", {"items": items, "form": form})


@login_required
def download_user_file(request, pk):
    item = get_object_or_404(UserFile, pk=pk)
    is_owner = item.owner_id == request.user.id
    is_admin = has_any_role(request.user, UserProfile.ROLE_ADMIN)
    if not is_owner and not is_admin:
        messages.error(request, "Недостаточно прав для скачивания этого файла.")
        return redirect("core:my_files")

    return FileResponse(item.file.open("rb"), as_attachment=True, filename=item.filename)


@role_required(UserProfile.ROLE_ADMIN)
def staff_dashboard(request):
    role_rows = UserProfile.objects.values("role").annotate(total=Count("id")).order_by("role")
    role_stats = {row["role"]: row["total"] for row in role_rows}

    context = {
        "total_users": User.objects.count(),
        "total_news": News.objects.count(),
        "total_events": Event.objects.count(),
        "total_feedbacks": FeedbackMessage.objects.count(),
        "total_join_requests": CouncilJoinApplication.objects.count(),
        "total_files": UserFile.objects.count(),
        "role_stats": {
            "admin": role_stats.get(UserProfile.ROLE_ADMIN, 0),
            "manager": role_stats.get(UserProfile.ROLE_MANAGER, 0),
            "student": role_stats.get(UserProfile.ROLE_STUDENT, 0),
        },
        "latest_feedbacks": FeedbackMessage.objects.select_related("user")[:10],
        "latest_join_requests": CouncilJoinApplication.objects.select_related("user")[:10],
    }
    return render(request, "core/staff/dashboard.html", context)


@role_required(UserProfile.ROLE_ADMIN)
def staff_feedbacks(request):
    if request.method == "POST":
        item = get_object_or_404(FeedbackMessage, pk=request.POST.get("feedback_id"))
        new_status = request.POST.get("status", "")
        moderation_comment = request.POST.get("moderation_comment", "").strip()
        valid_statuses = {choice[0] for choice in FeedbackMessage.STATUS_CHOICES}

        if new_status in valid_statuses:
            item.status = new_status
            item.moderation_comment = moderation_comment
            item.save(update_fields=["status", "moderation_comment", "updated_at"])
            messages.success(request, "Статус обращения обновлен.")
        else:
            messages.error(request, "Передан некорректный статус обращения.")
        return redirect("core:staff_feedbacks")

    items = FeedbackMessage.objects.select_related("user")
    return render(request, "core/staff/feedbacks.html", {"items": items})


@role_required(UserProfile.ROLE_ADMIN)
def staff_join_requests(request):
    if request.method == "POST":
        item = get_object_or_404(CouncilJoinApplication, pk=request.POST.get("join_id"))
        new_status = request.POST.get("status", "")
        moderation_comment = request.POST.get("moderation_comment", "").strip()
        valid_statuses = {choice[0] for choice in CouncilJoinApplication.STATUS_CHOICES}

        if new_status in valid_statuses:
            item.status = new_status
            item.moderation_comment = moderation_comment
            item.save(update_fields=["status", "moderation_comment", "updated_at"])
            messages.success(request, "Статус заявки обновлен.")
        else:
            messages.error(request, "Передан некорректный статус заявки.")
        return redirect("core:staff_join_requests")

    items = CouncilJoinApplication.objects.select_related("user")
    return render(request, "core/staff/join_requests.html", {"items": items})


@role_required(UserProfile.ROLE_ADMIN)
def staff_users(request):
    if request.method == "POST":
        profile = get_object_or_404(UserProfile, pk=request.POST.get("profile_id"))
        new_role = request.POST.get("role", "")
        valid_roles = {choice[0] for choice in UserProfile.ROLE_CHOICES}

        if new_role in valid_roles:
            profile.role = new_role
            profile.save(update_fields=["role"])
            messages.success(request, f"Роль пользователя {profile.user.username} обновлена.")
        else:
            messages.error(request, "Передано недопустимое значение роли.")
        return redirect("core:staff_users")

    profiles = UserProfile.objects.select_related("user").order_by("user__username")
    return render(request, "core/staff/users.html", {"profiles": profiles})


@role_required(UserProfile.ROLE_ADMIN)
def staff_files(request):
    if request.method == "POST" and request.POST.get("action") == "delete":
        item = get_object_or_404(UserFile, pk=request.POST.get("file_id"))
        item.file.delete(save=False)
        item.delete()
        messages.success(request, "Файл удален из хранилища.")
        return redirect("core:staff_files")

    items = UserFile.objects.select_related("owner", "owner__profile")
    return render(request, "core/staff/files.html", {"items": items})


@role_required(UserProfile.ROLE_ADMIN)
def staff_reports(request):
    context = {
        "events_total": Event.objects.count(),
        "feedbacks_total": FeedbackMessage.objects.count(),
        "generated_at": timezone.now(),
    }
    return render(request, "core/staff/reports.html", context)


@role_required(UserProfile.ROLE_ADMIN)
def export_events_docx(request):
    try:
        from docx import Document as DocxDocument
    except ImportError:
        messages.error(request, "Для экспорта DOCX установите библиотеку python-docx.")
        return redirect("core:staff_reports")

    events = Event.objects.order_by("start_at")

    doc = DocxDocument()
    doc.add_heading("Отчет по мероприятиям студсовета", level=1)
    doc.add_paragraph(f"Дата формирования: {timezone.localtime(timezone.now()).strftime('%d.%m.%Y %H:%M')}")

    for event in events:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.add_run(event.title).bold = True
        paragraph.add_run(f"\nДата: {timezone.localtime(event.start_at).strftime('%d.%m.%Y %H:%M')}")
        paragraph.add_run(f"\nМесто: {event.location}")
        paragraph.add_run(f"\nОписание: {event.short_description}")

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    filename = f"events_report_{timezone.localdate().isoformat()}.docx"
    return FileResponse(output, as_attachment=True, filename=filename)


@role_required(UserProfile.ROLE_ADMIN)
def export_feedback_xlsx(request):
    try:
        from openpyxl import Workbook
    except ImportError:
        messages.error(request, "Для экспорта XLSX установите библиотеку openpyxl.")
        return redirect("core:staff_reports")

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Feedback"
    sheet.append(["ID", "Дата", "Имя", "Email", "Тема", "Статус", "Комментарий модератора"])

    for item in FeedbackMessage.objects.order_by("-created_at"):
        sheet.append(
            [
                item.id,
                timezone.localtime(item.created_at).strftime("%d.%m.%Y %H:%M"),
                item.name,
                item.email,
                item.subject,
                item.get_status_display(),
                item.moderation_comment,
            ]
        )

    output = BytesIO()
    workbook.save(output)
    output.seek(0)

    filename = f"feedback_report_{timezone.localdate().isoformat()}.xlsx"
    return FileResponse(output, as_attachment=True, filename=filename)

