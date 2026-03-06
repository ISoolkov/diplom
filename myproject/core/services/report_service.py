from io import BytesIO

from django.utils import timezone

from .exceptions import ServiceDependencyError


def build_events_docx(events):
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise ServiceDependencyError("Для экспорта DOCX установите библиотеку python-docx.") from exc

    doc = DocxDocument()
    doc.add_heading("Отчет по мероприятиям студсовета", level=1)
    doc.add_paragraph(f"Дата формирования: {timezone.localtime(timezone.now()).strftime('%d.%m.%Y %H:%M')}")

    for event in events:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.add_run(event.title).bold = True
        paragraph.add_run(f"\nДата: {timezone.localtime(event.start_at).strftime('%d.%m.%Y %H:%M')}")
        paragraph.add_run(f"\nМесто: {event.location}")
        paragraph.add_run(f"\nОписание: {event.short_description}")

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def build_feedback_xlsx(feedback_items):
    try:
        from openpyxl import Workbook
    except ImportError as exc:
        raise ServiceDependencyError("Для экспорта XLSX установите библиотеку openpyxl.") from exc

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Feedback"
    sheet.append(["ID", "Дата", "Имя", "Email", "Тема", "Статус", "Комментарий модератора"])

    for item in feedback_items:
        sheet.append(
            [
                item.id,
                timezone.localtime(item.created_at).strftime("%d.%m.%Y %H:%M"),
                item.name,
                item.email,
                item.subject,
                item.get_status_display(),
                item.moderation_comment,
            ]
        )

    output = BytesIO()
    workbook.save(output)
    output.seek(0)
    return output
